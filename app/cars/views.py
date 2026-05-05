# cars/views.py

from django.core.files.base import ContentFile
from docx import Document
from django.shortcuts import render, redirect, get_object_or_404
from django.db.models import Sum
from .models import (
    Car,
    Sale,
    VehicleDocument,
    VehicleDocumentImage,
    AppConfiguration,
    IPVA,
    Contract,
    ContractImage,
    InvoiceDocument,
)

from .forms import (
    CarForm,
    SaleForm,
    IPVAForm,
    InvoiceDocumentForm,
    VehicleDocumentForm,
    ContractForm,
    AccessPasswordCreateForm,
    CredentialsUnlockForm,
    ItauCredentialsForm,
    ChangeAccessPasswordForm,
)

def car_update(request, pk):
    car = get_object_or_404(Car, pk=pk)

    form = CarForm(request.POST or None, request.FILES or None, instance=car)

    if form.is_valid():
        car = form.save()

        # Se o usuário editou o carro e mudou o status para vendido,
        # criamos uma venda automaticamente se ainda não existir.
        if car.status == 'vendido' and not car.sales.exists():
            Sale.objects.create(
                car=car,
                sale_price=car.sale_price
            )

        return redirect('cars_list')

    return render(request, 'cars/car_form.html', {
        'form': form,
        'editing': True,
        'car': car,
    })


def car_delete(request, pk):
    car = get_object_or_404(Car, pk=pk)

    if request.method == 'POST':
        # Se o carro já tem venda registrada, NÃO podemos apagar o carro,
        # porque isso destruiria o histórico da venda.
        # Então apenas garantimos que ele fique como vendido.
        if car.sales.exists():
            car.status = 'vendido'
            car.save()
        else:
            car.delete()

        return redirect('cars_list')

    return render(request, 'cars/car_delete.html', {
        'car': car,
    })

def dashboard(request):
    carros_estoque = Car.objects.filter(
        status='estoque',
        sales__isnull=True
    ).distinct()

    vendas = Sale.objects.select_related('car').all()

    context = {
        'carros_venda': carros_estoque.count(),
        'carros_vendidos': vendas.count(),
        'total_investido': carros_estoque.aggregate(Sum('purchase_price'))['purchase_price__sum'] or 0,
        'lucro_total': vendas.aggregate(Sum('profit'))['profit__sum'] or 0,
        'carros': carros_estoque.order_by('-id')[:5],
    }

    return render(request, 'cars/dashboard.html', context)

def cars_list(request):
    cars = Car.objects.filter(
        status='estoque',
        sales__isnull=True
    ).distinct().order_by('-id')

    return render(request, 'cars/cars_list.html', {'cars': cars})


def car_create(request):
    form = CarForm(request.POST or None, request.FILES or None)

    if form.is_valid():
        car = form.save(commit=False)
        car.status = 'estoque'
        car.save()
        return redirect('cars_list')

    return render(request, 'cars/car_form.html', {
        'form': form,
        'editing': False,
    })

def sold_cars_list(request):
    query = request.GET.get('q', '')

    sales = Sale.objects.select_related('car').all().order_by('-created_at')

    if query:
        sales = sales.filter(
            car__brand__icontains=query
        ) | sales.filter(
            car__model__icontains=query
        ) | sales.filter(
            car__plate__icontains=query
        )

    total_sales = sales.count()
    total_value = sales.aggregate(Sum('sale_price'))['sale_price__sum'] or 0
    total_profit = sales.aggregate(Sum('profit'))['profit__sum'] or 0

    context = {
        'sales': sales,
        'total_sales': total_sales,
        'total_value': total_value,
        'total_profit': total_profit,
        'query': query,
    }

    return render(request, 'cars/sold_cars.html', context)


def sale_create(request):
    form = SaleForm(request.POST or None)

    if form.is_valid():
        form.save()
        return redirect('sold_cars_list')

    return render(request, 'cars/sale_form.html', {'form': form})



def documentation_list(request):
    car_id = request.GET.get('car')

    cars = Car.objects.all().order_by('brand', 'model')
    documents = VehicleDocument.objects.select_related('car').prefetch_related('images').all().order_by('-created_at')

    if car_id:
        documents = documents.filter(car_id=car_id)

    context = {
        'cars': cars,
        'documents': documents,
        'selected_car': car_id,
    }

    return render(request, 'cars/documentation.html', context)


def document_create(request):
    form = VehicleDocumentForm(request.POST or None, request.FILES or None)

    if form.is_valid():
        document = form.save()

        for image in request.FILES.getlist('images'):
            VehicleDocumentImage.objects.create(
                document=document,
                image=image
            )

        return redirect('documentation_list')

    return render(request, 'cars/document_form.html', {
        'form': form,
    })

def inventory_view(request):
    cars = Car.objects.all().order_by('-id')

    cars_in_stock = Car.objects.filter(
        status='estoque',
        sales__isnull=True
    ).distinct().order_by('-id')

    cars_for_sale = cars_in_stock

    total_inventory = cars_in_stock.count()
    total_investment = cars_in_stock.aggregate(Sum('purchase_price'))['purchase_price__sum'] or 0
    estimated_value = cars_in_stock.aggregate(Sum('fipe_current_price'))['fipe_current_price__sum'] or 0

    estimated_margin = estimated_value - total_investment

    if total_investment > 0:
        margin_percentage = (estimated_margin / total_investment) * 100
    else:
        margin_percentage = 0

    context = {
        'cars': cars_in_stock,
        'cars_in_stock': cars_in_stock,
        'cars_for_sale': cars_for_sale,
        'total_inventory': total_inventory,
        'total_investment': total_investment,
        'estimated_value': estimated_value,
        'estimated_margin': estimated_margin,
        'margin_percentage': margin_percentage,
    }

    return render(request, 'cars/inventory.html', context)

def financial_view(request):
    active_tab = request.GET.get('tab', 'ipva')

    # IPVA
    ipvas = IPVA.objects.select_related('car').all().order_by('-year', '-created_at')
    ipva_form = IPVAForm()

    # Contratos
    selected_car = request.GET.get('car', '')
    cars = Car.objects.all().order_by('brand', 'model', 'year')

    contracts = Contract.objects.select_related('car').prefetch_related('images').all().order_by('-created_at')

    if selected_car:
        contracts = contracts.filter(car_id=selected_car)

    contract_form = ContractForm()

    # Notas fiscais
    invoice_selected_car = request.GET.get('invoice_car', '')

    invoice_documents = InvoiceDocument.objects.select_related('car').all().order_by('-created_at')

    if invoice_selected_car:
        invoice_documents = invoice_documents.filter(car_id=invoice_selected_car)

    invoice_form = InvoiceDocumentForm()

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'create_ipva':
            active_tab = 'ipva'
            ipva_form = IPVAForm(request.POST)

            if ipva_form.is_valid():
                ipva_form.save()
                return redirect('/financeiro/?tab=ipva')

        elif action == 'create_contract':
            active_tab = 'contracts'
            contract_form = ContractForm(request.POST, request.FILES)

            if contract_form.is_valid():
                contract = contract_form.save()

                for image in request.FILES.getlist('images'):
                    ContractImage.objects.create(
                        contract=contract,
                        image=image
                    )

                return redirect('/financeiro/?tab=contracts')

        elif action == 'create_invoice':
            active_tab = 'invoices'
            invoice_form = InvoiceDocumentForm(request.POST, request.FILES)

            if invoice_form.is_valid():
                car = invoice_form.cleaned_data['car']
                title = invoice_form.cleaned_data['title']
                typed_text = invoice_form.cleaned_data.get('typed_text')
                uploaded_docx = invoice_form.cleaned_data.get('uploaded_docx')

                invoice = InvoiceDocument(
                    car=car,
                    title=title,
                )

                if uploaded_docx:
                    invoice.file_type = 'docx'
                    invoice.file = uploaded_docx
                    invoice.save()
                else:
                    invoice.file_type = 'txt'
                    invoice.text_content = typed_text
                    invoice.save()

                    filename = f'nota_fiscal_{invoice.id}.txt'
                    invoice.file.save(
                        filename,
                        ContentFile(typed_text.encode('utf-8')),
                        save=True
                    )

                return redirect('/financeiro/?tab=invoices')

    context = {
        'active_tab': active_tab,

        # IPVA
        'ipvas': ipvas,
        'ipva_form': ipva_form,

        # Contratos
        'cars': cars,
        'contracts': contracts,
        'contract_form': contract_form,
        'selected_car': selected_car,

        # Notas fiscais
        'invoice_documents': invoice_documents,
        'invoice_form': invoice_form,
        'invoice_selected_car': invoice_selected_car,
    }

    return render(request, 'cars/financial.html', context)


def invoice_document_detail(request, pk):
    invoice = get_object_or_404(InvoiceDocument, pk=pk)

    readable_content = ''

    if invoice.file_type == 'txt':
        if invoice.text_content:
            readable_content = invoice.text_content
        elif invoice.file:
            with invoice.file.open('rb') as f:
                readable_content = f.read().decode('utf-8', errors='replace')

    elif invoice.file_type == 'docx':
        if invoice.file:
            document = Document(invoice.file.path)

            paragraphs = []

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()

                if text:
                    paragraphs.append(text)

            readable_content = '\n\n'.join(paragraphs)

    return render(request, 'cars/invoice_document_detail.html', {
        'invoice': invoice,
        'readable_content': readable_content,
    })

def settings_view(request):
    config = AppConfiguration.load()

    active_tab = request.GET.get('tab', 'credentials')
    unlocked = request.session.get('itau_credentials_unlocked', False)

    access_form = AccessPasswordCreateForm()
    unlock_form = CredentialsUnlockForm()
    credentials_form = ItauCredentialsForm()
    change_password_form = ChangeAccessPasswordForm()

    message = ''
    error = ''

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'create_access_password':
            access_form = AccessPasswordCreateForm(request.POST)

            if access_form.is_valid():
                config.set_access_password(access_form.cleaned_data['new_password'])
                config.save()

                request.session['itau_credentials_unlocked'] = True
                return redirect('/configuracoes/?tab=credentials')

        elif action == 'unlock_credentials':
            unlock_form = CredentialsUnlockForm(request.POST)

            if unlock_form.is_valid():
                password = unlock_form.cleaned_data['access_password']

                if config.check_access_password(password):
                    request.session['itau_credentials_unlocked'] = True
                    return redirect('/configuracoes/?tab=credentials')

                error = 'Senha incorreta.'

        elif action == 'save_itau_credentials':
            credentials_form = ItauCredentialsForm(request.POST)

            if credentials_form.is_valid():
                config.set_itau_credentials(
                    credentials_form.cleaned_data['itau_email'],
                    credentials_form.cleaned_data['itau_password']
                )
                config.save()

                request.session['itau_credentials_unlocked'] = True
                return redirect('/configuracoes/?tab=credentials')

        elif action == 'lock_credentials':
            request.session['itau_credentials_unlocked'] = False
            return redirect('/configuracoes/?tab=credentials')

        elif action == 'change_access_password':
            active_tab = 'security'
            change_password_form = ChangeAccessPasswordForm(request.POST)

            if change_password_form.is_valid():
                current_password = change_password_form.cleaned_data['current_password']
                new_password = change_password_form.cleaned_data['new_password']

                if config.has_access_password() and not config.check_access_password(current_password):
                    error = 'Senha atual incorreta.'
                else:
                    config.set_access_password(new_password)
                    config.save()

                    request.session['itau_credentials_unlocked'] = False
                    message = 'Senha alterada com sucesso.'

    decrypted_email = ''
    decrypted_password = ''

    if unlocked and config.has_itau_credentials():
        try:
            decrypted_email = config.get_itau_email()
            decrypted_password = config.get_itau_password()
        except Exception:
            error = 'Não foi possível descriptografar as credenciais. Verifique a chave de criptografia.'

    context = {
        'active_tab': active_tab,
        'config': config,
        'unlocked': unlocked,
        'has_password': config.has_access_password(),
        'has_credentials': config.has_itau_credentials(),

        'access_form': access_form,
        'unlock_form': unlock_form,
        'credentials_form': credentials_form,
        'change_password_form': change_password_form,

        'decrypted_email': decrypted_email,
        'decrypted_password': decrypted_password,

        'message': message,
        'error': error,
    }

    return render(request, 'cars/settings.html', context)
